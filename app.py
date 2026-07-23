"""
LHP Kegiatan Positif — Flask web app
Google OAuth login, self-register (3 free tokens), weekly regen, Midtrans topup.
"""
import os, re, shutil, uuid, json, hashlib, hmac
import logging, traceback, time
from datetime import datetime, timezone, timedelta
from functools import wraps
from flask import (Flask, request, jsonify, send_file,
                   render_template, after_this_request,
                   session, redirect, url_for)
from werkzeug.utils import secure_filename

try:
    import pillow_heif
    pillow_heif.register_heif_opener()
    _HEIC_OK = True
except Exception:
    _HEIC_OK = False

BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_FILE = os.path.join(BASE_DIR, "template_lhp.docx")
XLSX_FILE     = os.path.join(BASE_DIR, "DATA_DANTON_DANKI.xlsx")
UPLOAD_FOLDER = os.path.join(BASE_DIR, "tmp")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ── Persistent data directory ─────────────────────────────────────────────────
DATA_DIR = (os.environ.get("DATA_DIR")
            or os.environ.get("RAILWAY_VOLUME_MOUNT_PATH")
            or BASE_DIR)
try:
    os.makedirs(DATA_DIR, exist_ok=True)
except Exception:
    DATA_DIR = BASE_DIR

USERS_FILE    = os.path.join(DATA_DIR, "users.json")
ORDERS_FILE   = os.path.join(DATA_DIR, "orders.json")
_LEGACY_USERS = os.path.join(BASE_DIR, "users.json")
STORAGE_IS_PERSISTENT = os.path.abspath(DATA_DIR) != os.path.abspath(BASE_DIR)

if not os.path.exists(USERS_FILE) and os.path.exists(_LEGACY_USERS) \
        and os.path.abspath(_LEGACY_USERS) != os.path.abspath(USERS_FILE):
    try:
        shutil.copy2(_LEGACY_USERS, USERS_FILE)
    except Exception:
        pass

import sys as _sys
print("=" * 60, file=_sys.stderr)
print(f"[STORAGE] DATA_DIR  = {DATA_DIR}", file=_sys.stderr)
print(f"[STORAGE] PERSISTENT = {STORAGE_IS_PERSISTENT}", file=_sys.stderr)
print("=" * 60, file=_sys.stderr)
_sys.stderr.flush()

# ── Config ────────────────────────────────────────────────────────────────────
ADMIN_USERNAME      = os.environ.get("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD      = os.environ.get("ADMIN_PASSWORD", "admin123")
TOKENS_NEW_USER     = 3          # token gratis saat daftar
TOKENS_WEEKLY_REGEN = 1          # token regen per minggu (hanya jika token == 0)
TOKENS_PER_DOC      = 1

# Google OAuth
GOOGLE_CLIENT_ID     = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")
GOOGLE_REDIRECT_URI  = os.environ.get("GOOGLE_REDIRECT_URI",
                                      "https://yourapp.railway.app/auth/google/callback")

# Midtrans
MIDTRANS_SERVER_KEY  = os.environ.get("MIDTRANS_SERVER_KEY", "")
MIDTRANS_CLIENT_KEY  = os.environ.get("MIDTRANS_CLIENT_KEY", "")
MIDTRANS_IS_PROD     = os.environ.get("MIDTRANS_ENV", "sandbox") == "production"
MIDTRANS_BASE_URL    = ("https://app.midtrans.com" if MIDTRANS_IS_PROD
                        else "https://app.sandbox.midtrans.com")
MIDTRANS_API_URL     = ("https://api.midtrans.com" if MIDTRANS_IS_PROD
                        else "https://api.sandbox.midtrans.com")

# Token packages
TOKEN_PACKAGES = [
    {"id": "pkg_1",  "tokens": 1,  "price": 5000,  "label": "1 Token",   "desc": "Rp5.000"},
    {"id": "pkg_5",  "tokens": 5,  "price": 25000, "label": "5 Token",   "desc": "Rp25.000"},
    {"id": "pkg_10", "tokens": 10, "price": 40000, "label": "10 Token",  "desc": "Rp40.000"},
]
PKG_MAP = {p["id"]: p for p in TOKEN_PACKAGES}

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "lhp-akpol-secret-2026-xK9mP")
app.config["MAX_CONTENT_LENGTH"] = 64 * 1024 * 1024
logging.basicConfig(level=logging.INFO)

@app.errorhandler(413)
def _too_large(e):
    return jsonify({'error': 'Total ukuran foto terlalu besar. Coba kurangi jumlah/ukuran foto.'}), 413

@app.errorhandler(Exception)
def _unhandled(e):
    from werkzeug.exceptions import HTTPException
    if isinstance(e, HTTPException):
        return e
    app.logger.error("UNHANDLED ERROR:\n%s", traceback.format_exc())
    return jsonify({'error': 'Kesalahan server. Coba lagi.'}), 500

# ══════════════════════════════════════════════════════════════════════════════
# User Store
# ══════════════════════════════════════════════════════════════════════════════

def _load_users():
    if not os.path.exists(USERS_FILE):
        return {}
    try:
        with open(USERS_FILE) as f:
            return json.load(f)
    except Exception:
        return {}

def _save_users(users):
    with open(USERS_FILE, 'w') as f:
        json.dump(users, f, indent=2, ensure_ascii=False)

def get_user(username):
    """Ambil user berdasarkan username (lowercase). Returns user dict or None."""
    return _load_users().get(username.lower())

def get_user_by_google_id(google_id):
    """Cari user berdasarkan google_id yang tersimpan."""
    users = _load_users()
    for u in users.values():
        if u.get('google_id') == str(google_id):
            return u
    return None

def is_google_id_used(google_id):
    """Cek apakah google_id sudah pernah dipakai daftar."""
    return get_user_by_google_id(google_id) is not None

def create_user(username, password, name, google_id, google_email, google_picture=''):
    """Buat akun baru setelah verifikasi Google. Returns (user, error)."""
    from werkzeug.security import generate_password_hash
    users = _load_users()
    key   = username.lower().strip()
    # Validasi username
    if not key:
        return None, 'Username tidak boleh kosong'
    if not re.match(r'^[a-z0-9._-]{3,30}$', key):
        return None, 'Username hanya boleh huruf kecil, angka, titik, underscore, 3-30 karakter'
    if key in users:
        return None, 'Username sudah digunakan, pilih username lain'
    # Cek google_id belum pernah dipakai
    if is_google_id_used(google_id):
        return None, 'Akun Google ini sudah pernah digunakan untuk mendaftar'
    now = datetime.now(timezone.utc).isoformat()
    users[key] = {
        'uid':            key,
        'username':       key,
        'name':           name.strip(),
        'password':       generate_password_hash(password),
        'google_id':      str(google_id),
        'google_email':   google_email,
        'picture':        google_picture,
        'tokens':         TOKENS_NEW_USER,
        'created_at':     now,
        'last_regen':     now,
    }
    _save_users(users)
    app.logger.info("New user registered: %s (google: %s)", key, google_email)
    return users[key], None

def delete_user(username):
    users = _load_users()
    key   = username.lower()
    if key not in users:
        return False
    del users[key]
    _save_users(users)
    return True

def use_token(username):
    users = _load_users()
    key   = username.lower()
    if key not in users:
        return False
    if users[key].get('tokens', 0) < TOKENS_PER_DOC:
        return False
    users[key]['tokens'] -= TOKENS_PER_DOC
    _save_users(users)
    return True

def add_tokens(username, amount):
    users = _load_users()
    key   = username.lower()
    if key not in users:
        return False
    users[key]['tokens'] = users[key].get('tokens', 0) + amount
    _save_users(users)
    return True

def try_weekly_regen(username):
    """Regenerasi 1 token/minggu HANYA jika token == 0."""
    users = _load_users()
    key   = username.lower()
    if key not in users:
        return False, None
    user = users[key]
    if user.get('tokens', 0) > 0:
        return False, None

    last_str = user.get('last_regen', user.get('created_at'))
    try:
        last_dt = datetime.fromisoformat(last_str)
        if last_dt.tzinfo is None:
            last_dt = last_dt.replace(tzinfo=timezone.utc)
    except Exception:
        last_dt = datetime.now(timezone.utc) - timedelta(weeks=2)

    now        = datetime.now(timezone.utc)
    next_regen = last_dt + timedelta(weeks=1)

    if now >= next_regen:
        users[key]['tokens']     = TOKENS_WEEKLY_REGEN
        users[key]['last_regen'] = now.isoformat()
        _save_users(users)
        return True, (now + timedelta(weeks=1)).isoformat()

    return False, next_regen.isoformat()

# ══════════════════════════════════════════════════════════════════════════════
# Order Store (Midtrans)
# ══════════════════════════════════════════════════════════════════════════════

def _load_orders():
    if not os.path.exists(ORDERS_FILE):
        return {}
    try:
        with open(ORDERS_FILE) as f:
            return json.load(f)
    except Exception:
        return {}

def _save_orders(orders):
    with open(ORDERS_FILE, 'w') as f:
        json.dump(orders, f, indent=2, ensure_ascii=False)

def create_order(uid, pkg_id):
    pkg    = PKG_MAP.get(pkg_id)
    if not pkg:
        return None, "Paket tidak valid"
    order_id = f"LHP-{uid[:8]}-{uuid.uuid4().hex[:8].upper()}"
    orders   = _load_orders()
    orders[order_id] = {
        'order_id':  order_id,
        'uid':       uid,
        'pkg_id':    pkg_id,
        'tokens':    pkg['tokens'],
        'amount':    pkg['price'],
        'status':    'pending',
        'created_at': datetime.now(timezone.utc).isoformat(),
    }
    _save_orders(orders)
    return order_id, None

def get_order(order_id):
    return _load_orders().get(order_id)

def complete_order(order_id):
    orders = _load_orders()
    if order_id not in orders:
        return False
    if orders[order_id]['status'] == 'paid':
        return True   # idempotent
    orders[order_id]['status']  = 'paid'
    orders[order_id]['paid_at'] = datetime.now(timezone.utc).isoformat()
    _save_orders(orders)
    add_tokens(orders[order_id]['username'], orders[order_id]['tokens'])
    return True

# ══════════════════════════════════════════════════════════════════════════════
# Auth decorators
# ══════════════════════════════════════════════════════════════════════════════

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'uid' not in session:
            if request.is_json or request.path.startswith('/api/'):
                return jsonify({'error': 'Login diperlukan', 'redirect': '/login'}), 401
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'uid' not in session:
            if request.is_json or request.path.startswith('/api/'):
                return jsonify({'error': 'Sesi habis. Silakan login ulang.', 'redirect': '/login-admin'}), 401
            return redirect(url_for('login_admin'))
        if session.get('role') != 'admin':
            if request.is_json or request.path.startswith('/api/'):
                return jsonify({'error': 'Akses ditolak.'}), 403
            return redirect(url_for('login_admin'))
        return f(*args, **kwargs)
    return decorated

# ══════════════════════════════════════════════════════════════════════════════
# Google OAuth helpers
# ══════════════════════════════════════════════════════════════════════════════

def _google_auth_url():
    import urllib.parse
    params = {
        'client_id':     GOOGLE_CLIENT_ID,
        'redirect_uri':  GOOGLE_REDIRECT_URI,
        'response_type': 'code',
        'scope':         'openid email profile',
        'access_type':   'online',
        'prompt':        'select_account',
    }
    return 'https://accounts.google.com/o/oauth2/v2/auth?' + urllib.parse.urlencode(params)

def _google_exchange_code(code):
    """Exchange auth code for id_token + access_token. Returns dict or raises."""
    import urllib.request, urllib.parse
    payload = urllib.parse.urlencode({
        'code':          code,
        'client_id':     GOOGLE_CLIENT_ID,
        'client_secret': GOOGLE_CLIENT_SECRET,
        'redirect_uri':  GOOGLE_REDIRECT_URI,
        'grant_type':    'authorization_code',
    }).encode()
    req  = urllib.request.Request('https://oauth2.googleapis.com/token',
                                  data=payload, method='POST')
    req.add_header('Content-Type', 'application/x-www-form-urlencoded')
    with urllib.request.urlopen(req, timeout=10) as r:
        return json.loads(r.read())

def _google_userinfo(access_token):
    import urllib.request
    req = urllib.request.Request('https://www.googleapis.com/oauth2/v2/userinfo')
    req.add_header('Authorization', f'Bearer {access_token}')
    with urllib.request.urlopen(req, timeout=10) as r:
        return json.loads(r.read())

# ══════════════════════════════════════════════════════════════════════════════
# Midtrans helpers
# ══════════════════════════════════════════════════════════════════════════════

def _midtrans_create_transaction(order_id, amount, name, email):
    """Buat Snap token via Midtrans Snap API."""
    import urllib.request, base64
    snap_url = ("https://app.midtrans.com/snap/v1/transactions"
                if MIDTRANS_IS_PROD
                else "https://app.sandbox.midtrans.com/snap/v1/transactions")
    auth  = base64.b64encode(f"{MIDTRANS_SERVER_KEY}:".encode()).decode()
    parts = name.strip().split(' ', 1)
    first_name = parts[0]
    last_name  = parts[1] if len(parts) > 1 else ''
    payload = json.dumps({
        "transaction_details": {
            "order_id":     order_id,
            "gross_amount": int(amount)
        },
        "customer_details": {
            "first_name": first_name,
            "last_name":  last_name,
            "email":      email,
        },
        "credit_card": {"secure": True},
    }).encode()
    req = urllib.request.Request(snap_url, data=payload, method='POST')
    req.add_header('Authorization', f'Basic {auth}')
    req.add_header('Content-Type',  'application/json')
    req.add_header('Accept',        'application/json')
    with urllib.request.urlopen(req, timeout=15) as r:
        return json.loads(r.read())

def _midtrans_verify_notif(notif: dict) -> bool:
    """Verify Midtrans notification signature."""
    if not MIDTRANS_SERVER_KEY:
        return True  # dev mode
    raw    = (notif.get('order_id','') +
              notif.get('status_code','') +
              notif.get('gross_amount','') +
              MIDTRANS_SERVER_KEY)
    sig    = hashlib.sha512(raw.encode()).hexdigest()
    return hmac.compare_digest(sig, notif.get('signature_key',''))

# ══════════════════════════════════════════════════════════════════════════════
# Utility helpers (carry over)
# ══════════════════════════════════════════════════════════════════════════════

def to_roman(n) -> str:
    try: n = int(n)
    except: return str(n)
    vals = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),(90,'XC'),
            (50,'L'),(40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    r = ''
    for v, s in vals:
        while n >= v: r += s; n -= v
    return r or str(n)

PANGKAT_SINGKAT = {
    'BHAYANGKARA TARUNA':     'BHATAR',
    'AJUN BRIGADIR TARUNA':   'ABRIGTAR',
    'BRIGADIR TARUNA':        'BRIGTAR',
    'BRIGADIR KEPALA TARUNA': 'BRIGKATAR',
}
def pangkat_singkat(p: str) -> str:
    return PANGKAT_SINGKAT.get(p.upper().strip(), p.upper().strip())

TINGKAT_CONFIG = {
    '1': {'kop': 'BATALYON TARUNA TK I/61/X',    'header': 'LAPORAN KEGIATAN TARUNA TK. I/61/X',
          'angkatan': 'ANGKATAN KE-61, BATALYON X',                  'tk_suffix': 'TK I/61/X'},
    '2': {'kop': 'BATALYON TARUNA TK II/60/MS',   'header': 'LAPORAN KEGIATAN TARUNA TK. II/60/MS',
          'angkatan': 'ANGKATAN KE-60, BATALYON MANGGALA SATYA',     'tk_suffix': 'TK II/60/MS'},
    '3': {'kop': 'BATALYON TARUNA TK III/59/BD',  'header': 'LAPORAN KEGIATAN TARUNA TK. III/59/BD',
          'angkatan': 'ANGKATAN KE-59, BATALYON BHAYANGKARA DHARMA', 'tk_suffix': 'TK III/59/BD'},
}

MONTHS_ID    = ['','JANUARI','FEBRUARI','MARET','APRIL','MEI','JUNI',
                'JULI','AGUSTUS','SEPTEMBER','OKTOBER','NOVEMBER','DESEMBER']
MONTHS_SHORT = {'jan':1,'feb':2,'mar':3,'apr':4,'mei':5,'jun':6,
                'jul':7,'agu':8,'sep':9,'okt':10,'nov':11,'des':12}

def parse_tanggal(s):
    s = s.strip(); hari = ''
    parts = s.split(',', 1)
    date_part = parts[1].strip() if len(parts) == 2 else s
    if len(parts) == 2: hari = parts[0].strip().upper()
    tokens = date_part.split()
    tgl_num   = tokens[0] if tokens else ''
    bulan_raw = tokens[1].lower().rstrip('.') if len(tokens) > 1 else ''
    tahun_str = tokens[2] if len(tokens) > 2 else ''
    idx = MONTHS_SHORT.get(bulan_raw[:3], 0)
    bulan_str = MONTHS_ID[idx] if idx else bulan_raw.upper()
    return hari, tgl_num, bulan_str, tahun_str

def parse_waktu(s):
    s = re.sub(r'\s*(WIB|WITA|WIT)\s*$', '', s.strip(), flags=re.IGNORECASE)
    return s.replace(':', '.').strip()

def _load_xlsx_for_tingkat(tingkat: str):
    sheet_map = {'1': 'TK I', '2': 'TK II', '3': 'TK III'}
    target_sheet = sheet_map.get(str(tingkat), 'TK II')
    try:
        import openpyxl
        wb = openpyxl.load_workbook(XLSX_FILE, data_only=True)
        ws = None
        for name in wb.sheetnames:
            if name.strip().upper() == target_sheet.upper():
                ws = wb[name]; break
        if ws is None:
            ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        if not rows: return []
        headers = [str(h).strip() if h else '' for h in rows[0]]
        return [{headers[i]: (str(rows[j][i]).strip() if rows[j][i] is not None else '')
                 for i in range(len(headers))}
                for j in range(1, len(rows)) if any(rows[j])]
    except Exception:
        return []

def _norm(jab):
    jab = jab.upper().strip()
    jab = re.sub(r'DANKITAR',  'DANKI TAR',  jab)
    jab = re.sub(r'DANTONTAR', 'DANTON TAR', jab)
    return re.sub(r'\s+', ' ', jab)

def lookup_danton(peleton, kompi, tingkat='2'):
    target = f"DANTON TAR {peleton}/{to_roman(kompi)}"
    for row in _load_xlsx_for_tingkat(tingkat):
        if _norm(row.get('JABATAN','')) == target:
            return {'Nama Danton': row.get('NAMA',''),
                    'Pangkat Danton': row.get('PANGKAT',''),
                    'NRP Danton': row.get('NRP','')}
    return None

def lookup_danki(kompi, tingkat='2'):
    target = f"DANKI TAR {to_roman(kompi)}"
    for row in _load_xlsx_for_tingkat(tingkat):
        if _norm(row.get('JABATAN','')) == target:
            return {'Nama Danki': row.get('NAMA',''),
                    'Pangkat Danki': row.get('PANGKAT',''),
                    'NRP Danki': row.get('NRP','')}
    return None

# ── Docx helpers (unchanged) ──────────────────────────────────────────────────

def _has_drawing(run):
    from docx.oxml.ns import qn
    el = run._element
    return (el.find(qn('w:drawing')) is not None or
            el.find('{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent') is not None)

def _replace_para(para, old, new):
    text_runs = [r for r in para.runs if not _has_drawing(r)]
    full = ''.join(r.text for r in text_runs)
    if old not in full: return False
    if text_runs:
        text_runs[0].text = full.replace(old, new)
        for r in text_runs[1:]: r.text = ''
    return True

def _all_paragraphs(doc):
    for p in doc.paragraphs: yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs: yield p

def _downscale_image(path, max_edge=1280, quality=82):
    try:
        from PIL import Image, ImageOps
        with Image.open(path) as im:
            im = ImageOps.exif_transpose(im)
            if im.mode not in ('RGB', 'L'):
                im = im.convert('RGB')
            w, h = im.size
            scale = min(1.0, max_edge / float(max(w, h)))
            if scale < 1.0:
                im = im.resize((max(1, int(w*scale)), max(1, int(h*scale))), Image.LANCZOS)
            out = path + '.small.jpg'
            im.save(out, 'JPEG', quality=quality, optimize=True)
            return out
    except Exception:
        return path

def _insert_images(doc, placeholder, image_paths):
    from docx.shared import Inches
    from lxml import etree
    target = None
    for para in _all_paragraphs(doc):
        if placeholder in ''.join(r.text for r in para.runs):
            target = para; break
    if not target: return
    for r in target.runs: r.text = ''
    _scaled_tmp = []
    for path in image_paths:
        if path and os.path.exists(path):
            small = _downscale_image(path)
            if small != path: _scaled_tmp.append(small)
            run = target.add_run()
            try: run.add_picture(small, width=Inches(2.8))
            except: run.text = f'[{os.path.basename(path)}]'
    for p in _scaled_tmp:
        try: os.remove(p)
        except: pass
    body = doc.element.body
    try: img_idx = list(body).index(target._element)
    except ValueError: return
    to_remove = []
    for el in list(body)[img_idx + 1:]:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag != 'p': break
        texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', etree.tostring(el, encoding='unicode'))
        if ''.join(texts).strip(): break
        to_remove.append(el)
    for el in to_remove: body.remove(el)

def _fix_signature_formatting(doc, pangkat_danton, nrp_danton,
                               nama_danki, pangkat_danki, nrp_danki,
                               pangkat_abbr='ABRIGTAR'):
    from lxml import etree
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    def wtag(n): return f'{{{W}}}{n}'
    def rm_italic(para):
        for rpr in para._element.iter(wtag('rPr')):
            for tag in (wtag('i'), wtag('iCs')):
                for el in list(rpr.findall(tag)): rpr.remove(el)
    def set_center(para):
        pPr = para._element.find(wtag('pPr'))
        if pPr is None:
            pPr = etree.SubElement(para._element, wtag('pPr'))
            para._element.insert(0, pPr)
        jc = pPr.find(wtag('jc'))
        if jc is None: jc = etree.SubElement(pPr, wtag('jc'))
        jc.set(wtag('val'), 'center')
    for para in _all_paragraphs(doc):
        full = ''.join(r.text for r in para.runs if not _has_drawing(r))
        if ('(Nama lengkap dan gelar dankitar)' in full or
            (nama_danki and nama_danki in full and pangkat_abbr not in full
             and 'NRP' not in full and '(Nama lengkap dan gelar dantontar)' not in full)):
            rm_italic(para)
        elif ('(Pangkat danki)' in full or '(NRP Danki)' in full or
              (pangkat_danki and pangkat_danki in full and nrp_danki and
               nrp_danki in full and pangkat_abbr not in full)):
            rm_italic(para)
        if ('(Pangkat danton)' in full or '(NRP Danton)' in full or
            (pangkat_danton and pangkat_danton in full and
             nrp_danton and nrp_danton in full and pangkat_abbr in full)):
            set_center(para)

def fill_template(data, image_paths, output_path):
    from docx import Document
    shutil.copy(TEMPLATE_FILE, output_path)
    doc = Document(output_path)

    nama           = data['Nama']
    no_ak          = data['No Ak']
    pangkat        = data['Pangkat'].upper()
    pangkat_abbr   = pangkat_singkat(pangkat)
    peleton        = data['Peleton']
    kompi_roman    = to_roman(data['Kompi'])
    nama_kegiatan  = data['Nama Kegiatan']
    tanggal_raw    = data['Tanggal Kegiatan']
    waktu_raw      = data['Waktu Kegiatan']
    tempat         = data['Tempat Kegiatan']
    nama_danton    = data['Nama Danton']
    pangkat_danton = data['Pangkat Danton']
    nrp_danton     = data['NRP Danton']
    nama_danki     = data['Nama Danki']
    pangkat_danki  = data['Pangkat Danki']
    nrp_danki      = data['NRP Danki']

    tingkat      = str(data.get('Tingkat', '2'))
    tk_cfg       = TINGKAT_CONFIG.get(tingkat, TINGKAT_CONFIG['2'])
    kop_baru     = tk_cfg['kop']
    header_baru  = tk_cfg['header']
    angkatan_str = tk_cfg['angkatan']
    tk_suffix    = tk_cfg['tk_suffix']

    hari, tgl_num, bulan_str, tahun_str = parse_tanggal(tanggal_raw)
    waktu_clean = parse_waktu(waktu_raw)

    uraian = (
        f"--------PADA HARI {hari} TANGGAL {tgl_num} BULAN {bulan_str} "
        f"TAHUN {tahun_str} PUKUL {waktu_clean} WIB, SAYA {nama} "
        f"TARUNA AKPOL, PANGKAT {pangkat}, NO AKADEMI {no_ak}, "
        f"{angkatan_str}, "
        f"TELAH MELAKSANAKAN KEGIATAN POSITIF BERUPA {nama_kegiatan.upper()}.-"
    )

    simple = {
        'BATALYON TARUNA TK I/60/MS':             kop_baru,
        'LAPORAN KEGIATAN TARUNA TK. I/60/MS':    header_baru,
        'DANTONTAR 1 KOMPI III':                  f'DANTONTAR {peleton} KOMPI {kompi_roman}',
        'TK I/60/MS YANG MEMBUAT LAPORAN':        f'{tk_suffix} YANG MEMBUAT LAPORAN',
        '(No. Ak. Panjang)':                      no_ak,
        '(Nama lengkap taruna)':                  nama,
        'KOMPI III':                              f'KOMPI {kompi_roman}',
        'PLETON 1':                               f'PLETON {peleton}',
        '(Judul Kegiatan)':                       nama_kegiatan,
        '(Hari, Tanggal, pukul)':                 f'{tanggal_raw}, {waktu_raw}',
        '(Lokasi pelaksanaan kegiatan, lengkap)':  tempat,
        '(Tempat)':                               'Semarang',
        '(Tanggal Bulan Tahun)':                  tanggal_raw,
        '(Nama lengkap dan gelar dantontar)':     nama_danton,
        '(Pangkat danton)':                       pangkat_danton,
        '(NRP Danton)':                           nrp_danton,
        'ABRIGTAR':                               pangkat_abbr,
        '(No Ak ttd)':                            no_ak,
        '(Nama Lengkap)':                         nama,
        'DANKITAR III':                           f'DANKITAR {kompi_roman}',
        '(Nama lengkap dan gelar dankitar)':      nama_danki,
        '(Pangkat danki)':                        pangkat_danki,
        '(NRP Danki)':                            nrp_danki,
        '(tempat ttd)':                           'Semarang',
        '(tanggal buat laporan ttd)':             tanggal_raw,
        '(Nama Lengkap ttd)':                     nama,
        '(BRIGTARakhir)':                         pangkat_abbr,
        '(No. Ak. Panjang ttd)':                  no_ak,
    }

    for para in _all_paragraphs(doc):
        full = ''.join(r.text for r in para.runs if not _has_drawing(r))
        if '--------PADA HARI' in full:
            text_runs = [r for r in para.runs if not _has_drawing(r)]
            if text_runs:
                text_runs[0].text = uraian
                for r in text_runs[1:]: r.text = ''
            continue
        for old, new in simple.items():
            _replace_para(para, old, new)
        _replace_para(para, 'TK I/60/MS', tk_suffix)

    valid_imgs = [p for p in image_paths if p and os.path.exists(p)]
    if valid_imgs:
        _insert_images(doc, '(Dokumentasi)', valid_imgs)

    _fix_signature_formatting(doc, pangkat_danton, nrp_danton,
                               nama_danki, pangkat_danki, nrp_danki,
                               pangkat_abbr)
    doc.save(output_path)

# ══════════════════════════════════════════════════════════════════════════════
# Routes — Auth
# Alur:
#   DAFTAR : /register → Google OAuth → /auth/google/callback → /register/form
#            → POST /api/register → login otomatis → /
#   LOGIN  : /login → POST username+password → /
#   ADMIN  : /login-admin → POST username+password → /admin
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def index():
    if 'uid' not in session:
        return redirect(url_for('login'))
    if session.get('role') == 'admin':
        return redirect(url_for('admin_panel'))
    user = get_user(session['uid'])
    if not user:
        session.clear()
        return redirect(url_for('login'))
    try_weekly_regen(session['uid'])
    user = get_user(session['uid'])
    return render_template('index.html',
                           user_name=user['name'],
                           user_tokens=user.get('tokens', 0),
                           user_picture=user.get('picture',''),
                           token_packages=TOKEN_PACKAGES,
                           midtrans_client_key=MIDTRANS_CLIENT_KEY)

@app.route('/admin/login', methods=['GET'])
def admin_login_redirect():
    return redirect(url_for('login_admin'))

# ── Login (username + password) ───────────────────────────────────────────────

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'uid' in session:
        return redirect(url_for('index'))
    if request.method == 'GET':
        return render_template('login.html',
                               google_configured=bool(GOOGLE_CLIENT_ID),
                               google_auth_url=_google_auth_url() if GOOGLE_CLIENT_ID else '')
    # POST — login dengan username + password
    data     = request.get_json() or {}
    username = data.get('username', '').strip().lower()
    password = data.get('password', '')
    if not username or not password:
        return jsonify({'error': 'Isi semua field'}), 400
    from werkzeug.security import check_password_hash
    user = get_user(username)
    if not user or not check_password_hash(user.get('password',''), password):
        return jsonify({'error': 'Username atau password salah'}), 401
    session['uid']     = user['username']
    session['role']    = 'user'
    session['name']    = user['name']
    session['picture'] = user.get('picture', '')
    return jsonify({'ok': True, 'redirect': '/'})

# ── Register step 1: mulai OAuth Google ───────────────────────────────────────

@app.route('/register')
def register():
    if 'uid' in session:
        return redirect(url_for('index'))
    if not GOOGLE_CLIENT_ID:
        return render_template('login.html', google_configured=False,
                               error='Google OAuth belum dikonfigurasi.')
    # Simpan state=register di session agar callback tahu ini alur daftar
    session['oauth_flow'] = 'register'
    return redirect(_google_auth_url())

# ── Register step 2: callback Google ─────────────────────────────────────────

@app.route('/auth/google/callback')
def auth_google_callback():
    error = request.args.get('error')
    if error:
        session.pop('oauth_flow', None)
        return redirect(url_for('login') + '?error=google_denied')
    code = request.args.get('code', '')
    if not code:
        session.pop('oauth_flow', None)
        return redirect(url_for('login') + '?error=no_code')
    try:
        token_data = _google_exchange_code(code)
        access_tok = token_data.get('access_token', '')
        userinfo   = _google_userinfo(access_tok)
        google_id  = str(userinfo['id'])
        email      = userinfo.get('email', '')
        name       = userinfo.get('name', email)
        picture    = userinfo.get('picture', '')
        flow       = session.pop('oauth_flow', 'register')

        if flow == 'register':
            # Cek apakah Google ID sudah pernah dipakai daftar
            existing = get_user_by_google_id(google_id)
            if existing:
                # Sudah punya akun — arahkan ke login dengan pesan
                return redirect(url_for('login') + '?error=google_already_registered')
            # Simpan data Google sementara di session, lanjut ke form registrasi
            session['pending_google'] = {
                'google_id':    google_id,
                'google_email': email,
                'picture':      picture,
                'suggested_name': name,
            }
            return redirect(url_for('register_form'))
        else:
            # Flow lain (tidak dipakai saat ini)
            return redirect(url_for('login'))
    except Exception:
        app.logger.error("Google OAuth error:\n%s", traceback.format_exc())
        session.pop('oauth_flow', None)
        return redirect(url_for('login') + '?error=oauth_failed')

# ── Register step 3: form isi nama/username/password ──────────────────────────

@app.route('/register/form', methods=['GET'])
def register_form():
    pending = session.get('pending_google')
    if not pending:
        return redirect(url_for('login') + '?error=session_expired')
    return render_template('register_form.html',
                           google_email=pending['google_email'],
                           suggested_name=pending['suggested_name'],
                           google_picture=pending['picture'])

@app.route('/api/register', methods=['POST'])
def api_register():
    pending = session.get('pending_google')
    if not pending:
        return jsonify({'error': 'Sesi habis. Ulangi proses pendaftaran.'}), 400
    data     = request.get_json() or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    name     = data.get('name', '').strip()
    if not username or not password or not name:
        return jsonify({'error': 'Semua field harus diisi'}), 400
    if len(password) < 6:
        return jsonify({'error': 'Password minimal 6 karakter'}), 400
    user, err = create_user(
        username=username,
        password=password,
        name=name,
        google_id=pending['google_id'],
        google_email=pending['google_email'],
        google_picture=pending['picture'],
    )
    if err:
        return jsonify({'error': err}), 400
    # Hapus pending data, langsung login
    session.pop('pending_google', None)
    session['uid']     = user['username']
    session['role']    = 'user'
    session['name']    = user['name']
    session['picture'] = user.get('picture', '')
    return jsonify({'ok': True, 'redirect': '/'})

# ── Admin login ───────────────────────────────────────────────────────────────

@app.route('/login-admin', methods=['GET', 'POST'])
def login_admin():
    if request.method == 'GET':
        return render_template('login_admin.html')
    data     = request.get_json() or {}
    username = data.get('username', '').strip().lower()
    password = data.get('password', '')
    if username == ADMIN_USERNAME.lower() and password == ADMIN_PASSWORD:
        session['uid']   = f'admin_{username}'
        session['role']  = 'admin'
        session['name']  = 'Admin'
        return jsonify({'ok': True, 'redirect': '/admin'})
    return jsonify({'error': 'Username atau password salah'}), 401

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/ads.txt')
def ads_txt():
    return send_file(os.path.join(BASE_DIR, 'static', 'ads.txt'), mimetype='text/plain')

@app.route('/robots.txt')
def robots_txt():
    return send_file(os.path.join(BASE_DIR, 'static', 'robots.txt'), mimetype='text/plain')

@app.route('/sitemap.xml')
def sitemap_xml():
    return send_file(os.path.join(BASE_DIR, 'static', 'sitemap.xml'), mimetype='application/xml')

@app.route('/google96f628609920d657.html')
def google_verify():
    return send_file(os.path.join(BASE_DIR, 'static', 'google96f628609920d657.html'), mimetype='text/html')

@app.route('/api/check-username')
def check_username_route():
    u   = request.args.get('u', '').strip().lower()
    ok  = bool(u) and re.match(r'^[a-z0-9._-]{3,30}$', u) is not None
    avail = ok and get_user(u) is None
    return jsonify({'available': avail})

# ══════════════════════════════════════════════════════════════════════════════
# Routes — Profile & Token info
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/api/token-balance')
@login_required
def token_balance():
    user = get_user(session['uid'])
    if not user:
        return jsonify({'tokens': 0})
    did_regen, next_regen = try_weekly_regen(session['uid'])
    user = get_user(session['uid'])
    return jsonify({
        'tokens':     user.get('tokens', 0),
        'did_regen':  did_regen,
        'next_regen': next_regen,
    })

# ══════════════════════════════════════════════════════════════════════════════
# Routes — Topup (Midtrans)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/api/topup/create', methods=['POST'])
@login_required
def topup_create():
    data   = request.get_json() or {}
    pkg_id = data.get('pkg_id', '')
    pkg    = PKG_MAP.get(pkg_id)
    if not pkg:
        return jsonify({'error': 'Paket tidak valid'}), 400

    uid  = session['uid']
    user = get_user(uid)
    if not user:
        return jsonify({'error': 'User tidak ditemukan'}), 404

    order_id, err = create_order(uid, pkg_id)
    if err:
        return jsonify({'error': err}), 400

    if not MIDTRANS_SERVER_KEY:
        # Dev mode — auto-complete without Midtrans
        complete_order(order_id)
        return jsonify({'ok': True, 'dev_mode': True,
                        'message': f'{pkg["tokens"]} token ditambahkan (dev mode)'})
    try:
        resp = _midtrans_create_transaction(
            order_id, pkg['price'],
            user.get('name', ''), user.get('email', ''))
        # Snap API returns 'redirect_url' for full-page redirect
        snap_token   = resp.get('token', '')
        redirect_url = resp.get('redirect_url', '')
        if not redirect_url and not snap_token:
            app.logger.error("Midtrans resp missing url/token: %s", resp)
            return jsonify({'error': 'Respons Midtrans tidak valid. Coba lagi.'}), 500
        return jsonify({'ok': True, 'payment_url': redirect_url,
                        'snap_token': snap_token, 'order_id': order_id})
    except Exception:
        app.logger.error("Midtrans error:\n%s", traceback.format_exc())
        return jsonify({'error': 'Gagal membuat transaksi. Coba lagi.'}), 500

@app.route('/api/topup/notification', methods=['POST'])
def topup_notification():
    """Midtrans webhook — called by Midtrans server after payment."""
    notif = request.get_json() or {}
    if not _midtrans_verify_notif(notif):
        app.logger.warning("Invalid Midtrans signature: %s", notif.get('order_id'))
        return jsonify({'ok': False}), 403

    order_id      = notif.get('order_id', '')
    txn_status    = notif.get('transaction_status', '')
    fraud_status  = notif.get('fraud_status', '')

    if txn_status in ('capture', 'settlement'):
        if fraud_status in ('accept', '') or txn_status == 'settlement':
            if complete_order(order_id):
                app.logger.info("Order paid: %s", order_id)
    elif txn_status in ('cancel', 'deny', 'expire', 'failure'):
        orders = _load_orders()
        if order_id in orders and orders[order_id]['status'] == 'pending':
            orders[order_id]['status'] = txn_status
            _save_orders(orders)

    return jsonify({'ok': True})

@app.route('/api/topup/status/<order_id>')
@login_required
def topup_status(order_id):
    order = get_order(order_id)
    if not order or order['uid'] != session['uid']:
        return jsonify({'error': 'Order tidak ditemukan'}), 404
    return jsonify({'status': order['status'], 'tokens': order['tokens']})

# ══════════════════════════════════════════════════════════════════════════════
# Routes — Admin panel
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/admin')
@admin_required
def admin_panel():
    users     = _load_users()
    user_list = sorted(users.values(), key=lambda u: u.get('created_at',''))
    return render_template('admin.html', users=user_list,
                           storage_persistent=STORAGE_IS_PERSISTENT,
                           data_dir=DATA_DIR)

@app.route('/api/admin/delete-user', methods=['POST'])
@admin_required
def admin_delete_user():
    data = request.get_json() or {}
    uid  = data.get('uid', '').strip()
    if not uid:
        return jsonify({'error': 'UID diperlukan'}), 400
    if delete_user(uid):
        return jsonify({'ok': True})
    return jsonify({'error': 'User tidak ditemukan'}), 404

@app.route('/api/admin/set-tokens', methods=['POST'])
@admin_required
def admin_set_tokens():
    data = request.get_json() or {}
    uid  = data.get('uid', '').strip()
    try:
        tokens = int(data.get('tokens', -1))
    except (ValueError, TypeError):
        return jsonify({'error': 'Jumlah token tidak valid'}), 400
    if tokens < 0 or tokens > 999:
        return jsonify({'error': 'Token 0–999'}), 400
    users = _load_users()
    if uid not in users:
        return jsonify({'error': 'User tidak ditemukan'}), 404
    users[uid]['tokens'] = tokens
    _save_users(users)
    return jsonify({'ok': True, 'tokens': tokens})

@app.route('/api/admin/reset-tokens', methods=['POST'])
@admin_required
def admin_reset_tokens():
    data  = request.get_json() or {}
    uid   = data.get('uid', '').strip()
    users = _load_users()
    if uid not in users:
        return jsonify({'error': 'User tidak ditemukan'}), 404
    users[uid]['tokens'] = TOKENS_NEW_USER
    _save_users(users)
    return jsonify({'ok': True, 'tokens': TOKENS_NEW_USER})

# ══════════════════════════════════════════════════════════════════════════════
# Routes — App (lookup + generate)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/api/lookup')
@login_required
def api_lookup():
    peleton = request.args.get('peleton', '').strip()
    kompi   = request.args.get('kompi', '').strip()
    tingkat = request.args.get('tingkat', '2').strip()
    if not peleton or not kompi:
        return jsonify({'error': 'peleton dan kompi diperlukan'}), 400
    if tingkat not in ('1', '2', '3'):
        tingkat = '2'
    danton = lookup_danton(peleton, kompi, tingkat)
    danki  = lookup_danki(kompi, tingkat)
    return jsonify({
        'danton': danton, 'danki': danki,
        'label':  f"DANTON TAR {peleton}/{to_roman(kompi)}  |  DANKI TAR {to_roman(kompi)}"
                  if danton and danki else None,
    })

@app.route('/api/generate', methods=['POST'])
@login_required
def api_generate():
    import tempfile
    uid  = session['uid']
    user = get_user(uid)
    if not user or user.get('tokens', 0) < TOKENS_PER_DOC:
        return jsonify({'error': 'Token habis. Topup token untuk melanjutkan.',
                        'no_token': True}), 402

    fields = ['Nama','No Ak','Pangkat','Tingkat','Peleton','Kompi',
              'Nama Danton','Pangkat Danton','NRP Danton',
              'Nama Danki','Pangkat Danki','NRP Danki',
              'Nama Kegiatan','Tanggal Kegiatan','Waktu Kegiatan','Tempat Kegiatan']
    data = {}
    for f in fields:
        val = request.form.get(f, '').strip()
        if not val:
            return jsonify({'error': f'Field "{f}" tidak boleh kosong'}), 400
        data[f] = val
    if data['Tingkat'] not in ('1', '2', '3'):
        return jsonify({'error': 'Tingkat tidak valid'}), 400

    image_paths, image_tmpfiles = [], []
    for i in range(1, 5):
        file = request.files.get(f'foto_{i}')
        if file and file.filename:
            ext = os.path.splitext(secure_filename(file.filename))[1] or '.jpg'
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext, dir=UPLOAD_FOLDER)
            file.save(tmp.name); tmp.close()
            image_paths.append(tmp.name)
            image_tmpfiles.append(tmp.name)

    out_tmp  = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=UPLOAD_FOLDER)
    out_path = out_tmp.name; out_tmp.close()
    out_name = f"LHP_{data['Nama'].replace(' ','_')}_{uuid.uuid4().hex[:6]}.docx"

    try:
        fill_template(data, image_paths, out_path)
    except Exception as e:
        app.logger.error("generate FAILED:\n%s", traceback.format_exc())
        for p in image_tmpfiles + [out_path]:
            try: os.remove(p)
            except: pass
        return jsonify({'error': f'Gagal membuat dokumen: {e}'}), 500
    finally:
        for p in image_tmpfiles:
            try: os.remove(p)
            except: pass

    use_token(uid)

    @after_this_request
    def cleanup(response):
        try: os.remove(out_path)
        except: pass
        return response

    return send_file(out_path, as_attachment=True, download_name=out_name,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
